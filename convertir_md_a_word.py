"""
Script para convertir el documento de Requerimientos Funcionales de Markdown a Word
con formato profesional utilizando python-docx.

Autor: Generado automáticamente
Fecha: 30 de enero de 2026
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def crear_estilos_documento(doc):
    """Crea estilos personalizados para el documento."""
    styles = doc.styles
    
    # Estilo para título principal
    if 'Título Principal' not in styles:
        style = styles.add_style('Título Principal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(24)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilo para subtítulos nivel 2
    if 'Subtítulo 2' not in styles:
        style = styles.add_style('Subtítulo 2', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    
    # Estilo para subtítulos nivel 3
    if 'Subtítulo 3' not in styles:
        style = styles.add_style('Subtítulo 3', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    
    # Estilo para subtítulos nivel 4
    if 'Subtítulo 4' not in styles:
        style = styles.add_style('Subtítulo 4', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(51, 51, 51)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

def procesar_linea_negrita(texto):
    """Procesa texto con formato **texto** para negrita."""
    partes = re.split(r'(\*\*.*?\*\*)', texto)
    return [(parte.strip('*'), parte.startswith('**')) for parte in partes if parte]

def agregar_parrafo_con_formato(doc, texto, estilo=None):
    """Agrega un párrafo procesando formato en negrita."""
    parrafo = doc.add_paragraph(style=estilo)
    
    # Buscar formato de negrita **texto**
    partes = re.split(r'(\*\*.*?\*\*)', texto)
    
    for parte in partes:
        if parte:
            if parte.startswith('**') and parte.endswith('**'):
                # Texto en negrita
                run = parrafo.add_run(parte.strip('*'))
                run.bold = True
            else:
                # Texto normal
                parrafo.add_run(parte)
    
    return parrafo

def procesar_tabla(lineas, idx, doc):
    """Procesa una tabla markdown y la convierte a tabla de Word."""
    tabla_lineas = []
    i = idx
    
    # Recopilar líneas de la tabla
    while i < len(lineas) and lineas[i].strip().startswith('|'):
        tabla_lineas.append(lineas[i].strip())
        i += 1
    
    if len(tabla_lineas) < 2:
        return i
    
    # Procesar encabezados
    encabezados = [col.strip() for col in tabla_lineas[0].split('|')[1:-1]]
    
    # Procesar filas de datos (saltar la línea de separación)
    filas_datos = []
    for linea in tabla_lineas[2:]:
        if linea.strip():
            celdas = [col.strip() for col in linea.split('|')[1:-1]]
            filas_datos.append(celdas)
    
    # Crear tabla en Word
    tabla = doc.add_table(rows=1 + len(filas_datos), cols=len(encabezados))
    tabla.style = 'Light Grid Accent 1'
    
    # Agregar encabezados
    for j, encabezado in enumerate(encabezados):
        celda = tabla.rows[0].cells[j]
        celda.text = encabezado
        # Formato de encabezado
        for paragraph in celda.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Agregar datos
    for i, fila in enumerate(filas_datos):
        for j, dato in enumerate(fila):
            tabla.rows[i + 1].cells[j].text = dato
    
    return idx + len(tabla_lineas)

def convertir_markdown_a_word(archivo_md, archivo_docx):
    """Convierte un archivo Markdown a Word con formato."""
    
    # Crear documento
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Crear estilos personalizados
    crear_estilos_documento(doc)
    
    # Leer archivo markdown
    with open(archivo_md, 'r', encoding='utf-8') as f:
        lineas = f.readlines()
    
    print(f"📄 Procesando {len(lineas)} líneas...")
    
    i = 0
    lista_activa = False
    lista_nivel = 0
    
    while i < len(lineas):
        linea = lineas[i].rstrip()
        
        # Línea vacía
        if not linea.strip():
            if not lista_activa:
                doc.add_paragraph()
            i += 1
            continue
        
        # Título nivel 1 (# Título)
        if linea.startswith('# ') and not linea.startswith('## '):
            titulo = linea[2:].strip()
            p = doc.add_paragraph(titulo, style='Título Principal')
            lista_activa = False
            i += 1
            continue
        
        # Título nivel 2 (## Título)
        if linea.startswith('## ') and not linea.startswith('### '):
            titulo = linea[3:].strip()
            p = doc.add_paragraph(titulo, style='Heading 1')
            lista_activa = False
            i += 1
            continue
        
        # Título nivel 3 (### Título)
        if linea.startswith('### ') and not linea.startswith('#### '):
            titulo = linea[4:].strip()
            p = doc.add_paragraph(titulo, style='Heading 2')
            lista_activa = False
            i += 1
            continue
        
        # Título nivel 4 (#### Título)
        if linea.startswith('#### '):
            titulo = linea[5:].strip()
            agregar_parrafo_con_formato(doc, titulo, 'Subtítulo 4')
            lista_activa = False
            i += 1
            continue
        
        # Línea horizontal (---)
        if linea.strip() in ['---', '***', '___']:
            doc.add_paragraph()
            i += 1
            continue
        
        # Tabla
        if linea.strip().startswith('|'):
            i = procesar_tabla(lineas, i, doc)
            lista_activa = False
            continue
        
        # Lista con viñetas (- item)
        if linea.strip().startswith('- ') or linea.strip().startswith('* '):
            texto = linea.strip()[2:].strip()
            p = agregar_parrafo_con_formato(doc, texto, 'List Bullet')
            lista_activa = True
            lista_nivel = 0
            i += 1
            continue
        
        # Lista numerada
        if re.match(r'^\d+\.\s', linea.strip()):
            texto = re.sub(r'^\d+\.\s', '', linea.strip())
            p = agregar_parrafo_con_formato(doc, texto, 'List Number')
            lista_activa = True
            lista_nivel = 0
            i += 1
            continue
        
        # Bloque de código (```...)
        if linea.strip().startswith('```'):
            codigo_lineas = []
            i += 1
            while i < len(lineas) and not lineas[i].strip().startswith('```'):
                codigo_lineas.append(lineas[i].rstrip())
                i += 1
            
            # Agregar código con formato
            if codigo_lineas:
                p = doc.add_paragraph('\n'.join(codigo_lineas))
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                # Fondo gris claro
                shading_elm = p._element.get_or_add_pPr()
            
            lista_activa = False
            i += 1
            continue
        
        # Texto normal
        if linea.strip():
            agregar_parrafo_con_formato(doc, linea.strip())
            lista_activa = False
        
        i += 1
    
    # Agregar tabla de contenidos (opcional)
    print("📋 Generando documento...")
    
    # Guardar documento
    doc.save(archivo_docx)
    print(f"✅ Documento guardado: {archivo_docx}")

def main():
    """Función principal."""
    archivo_entrada = 'REQUERIMIENTOS_FUNCIONALES.md'
    archivo_salida = 'REQUERIMIENTOS_FUNCIONALES.docx'
    
    print("=" * 60)
    print("🔄 CONVERSOR DE MARKDOWN A WORD")
    print("=" * 60)
    print(f"📥 Archivo de entrada: {archivo_entrada}")
    print(f"📤 Archivo de salida: {archivo_salida}")
    print()
    
    try:
        convertir_markdown_a_word(archivo_entrada, archivo_salida)
        print()
        print("=" * 60)
        print("✨ CONVERSIÓN COMPLETADA EXITOSAMENTE")
        print("=" * 60)
        print(f"📄 Puedes abrir el archivo: {archivo_salida}")
    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo {archivo_entrada}")
    except Exception as e:
        print(f"❌ Error durante la conversión: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
